"""
Generate EE627_Lecture_Notes.docx containing all updated Python codes.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Create a code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(8.5)
code_font.color.rgb = RGBColor(0, 0, 0)
code_fmt = code_style.paragraph_format
code_fmt.space_before = Pt(0)
code_fmt.space_after = Pt(0)
code_fmt.line_spacing = 1.0

# --- Title Page ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EE627 Microeconometrics')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Python Lecture Notes')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Instructor: Asst. Prof. Dr. Tiraphap Fakthong\nThammasat University')
run.font.size = Pt(14)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('This document contains the complete Python scripts for all 8 chapters,\n'
                    'designed to run on Google Colab.\n'
                    'Each script replicates the corresponding Stata do-file.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)

chapters = [
    ('Chapter 1', 'OLS Regression, Specification Analysis, and Prediction', 'Chap1_OLS.py'),
    ('Chapter 2', 'Monte Carlo Simulation Methods', 'Chap2_MonteCarlo.py'),
    ('Chapter 3', 'Heteroskedasticity, SUR, and Survey Data', 'Chap3_HetSUR.py'),
    ('Chapter 4', 'Instrumental Variables Estimation', 'Chap4_IV.py'),
    ('Chapter 5', 'Quantile Regression and Count Data Models', 'Chap5_Quantile.py'),
    ('Chapter 6', 'Panel Data: Linear and Nonlinear Models', 'Chap6_Panel.py'),
    ('Chapter 7', 'Regression Discontinuity Designs', 'Chap7_RD.py'),
    ('Chapter 8', 'Difference-in-Differences', 'Chap8_DID.py'),
]

for ch_num, ch_title, _ in chapters:
    p = doc.add_paragraph()
    run = p.add_run(f'{ch_num}: {ch_title}')
    run.font.size = Pt(12)

doc.add_page_break()

# --- Add each chapter ---
base_dir = os.path.dirname(os.path.abspath(__file__))

for ch_num, ch_title, filename in chapters:
    # Chapter heading
    doc.add_heading(f'{ch_num}: {ch_title}', level=1)

    # File info
    p = doc.add_paragraph()
    run = p.add_run(f'File: {filename}')
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # Read and add the Python code
    filepath = os.path.join(base_dir, filename)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            code = f.read()

        # Split into lines and add each line as a paragraph with code style
        lines = code.split('\n')
        for line in lines:
            p = doc.add_paragraph(style='CodeBlock')
            if line.strip() == '':
                # Empty line - add a space to preserve spacing
                p.add_run(' ')
            else:
                run = p.add_run(line)
                # Color comments differently
                stripped = line.lstrip()
                if stripped.startswith('#'):
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green for comments
                elif stripped.startswith('"""') or stripped.startswith("'''"):
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green for docstrings
                elif stripped.startswith('print(') or stripped.startswith('print ('):
                    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue for print

        print(f'  Added {ch_num} ({len(lines)} lines)')

    except FileNotFoundError:
        p = doc.add_paragraph()
        run = p.add_run(f'[File not found: {filename}]')
        run.font.color.rgb = RGBColor(255, 0, 0)
        print(f'  WARNING: {filename} not found')

    doc.add_page_break()

# --- Save ---
output_path = os.path.join(base_dir, 'EE627_Lecture_Notes.docx')
doc.save(output_path)
print(f'\nSaved: {output_path}')
